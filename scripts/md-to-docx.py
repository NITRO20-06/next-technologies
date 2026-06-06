#!/usr/bin/env python3
"""Convierte docs/TECNOLOGIAS.md a Word (.docx) con formato profesional."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT = ROOT / "docs" / "TECNOLOGIAS.md"
DEFAULT_OUTPUT = ROOT / "docs" / "TECNOLOGIAS.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text: str, base_bold: bool = False, base_italic: bool = False) -> None:
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()]).bold = base_bold
            paragraph.runs[-1].italic = base_italic

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = base_italic
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.bold = base_bold
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                label, url = link_match.groups()
                run = paragraph.add_run(f"{label} ({url})")
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                run.underline = True
        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        run.italic = base_italic


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s:-]+\|[\s|:-]+\|?$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    headers = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    table.style = "Table Grid"

    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        add_formatted_runs(p, header.replace("**", ""))
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
        set_cell_shading(cell, "E8EEF7")

    for row_idx, row in enumerate(body, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, value)
            for run in p.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], language: str = "") -> None:
    if language == "mermaid":
        p = doc.add_paragraph()
        run = p.add_run("Diagrama de arquitectura (resumen)")
        run.bold = True
        run.font.size = Pt(10)
        summary = [
            "Navegador: HTML, CSS, app.min.js, GSAP y Canvas de neblina.",
            "Entorno local: index.php, contacto.php, includes/ y data/contactos.log.",
            "Vercel: public/index.html, /api/contact, /api/csrf y build-vercel.mjs.",
            "Flujo: HTML → PHP (local) o build estático → API en producción.",
        ]
        for item in summary:
            bp = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(bp, item)
        return

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F1F5F9")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)


def convert_md_to_docx(input_path: Path, output_path: Path) -> None:
    lines = input_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lines = []
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            fence = stripped[3:].strip()
            in_code = True
            code_lang = fence
            code_lines = []
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if is_table_separator(stripped):
                i += 1
                continue
            table_rows.append(parse_table_row(stripped))
            i += 1
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue
        else:
            if table_rows:
                add_table(doc, table_rows)
                table_rows = []

        if stripped in ("---", "***"):
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, stripped[2:].strip())
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, numbered.group(2).strip())
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip("*").strip())
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            i += 1
            continue

        if stripped:
            p = doc.add_paragraph()
            add_formatted_runs(p, stripped)
        i += 1

    if table_rows:
        add_table(doc, table_rows)
    if in_code and code_lines:
        add_code_block(doc, code_lines, code_lang)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_INPUT
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT

    if not input_path.exists():
        print(f"No se encontró: {input_path}", file=sys.stderr)
        return 1

    convert_md_to_docx(input_path, output_path)
    print(f"Documento Word generado: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
